"""Generate the CLEAR consultancy report as a DOCX file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1A, 0x3A, 0x5C)   # headings
MID_BLUE   = RGBColor(0x23, 0x5A, 0x8C)   # subheadings
ACCENT     = RGBColor(0x00, 0x7A, 0xB5)   # table headers
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
BODY_GREY  = RGBColor(0x33, 0x33, 0x33)


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = DARK_BLUE if level == 1 else MID_BLUE
    run.font.bold = True
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_body(doc, text, bold=False, italic=False, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = BODY_GREY
    run.font.bold   = bold
    run.font.italic = italic
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = BODY_GREY
    p.paragraph_format.left_indent   = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after   = Pt(3)
    p.paragraph_format.space_before  = Pt(0)
    return p


def styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(9.5)
        set_cell_bg(cell, "007AB5")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size  = Pt(9.5)
            run.font.color.rgb = BODY_GREY
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths:
        for col_i, w in enumerate(col_widths):
            for cell in table.columns[col_i].cells:
                cell.width = Inches(w)

    doc.add_paragraph()
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),  "single")
    bottom.set(qn("w:sz"),   "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "007AB5")
    pb.append(bottom)
    pPr.append(pb)


# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

# ── Cover page ────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("CLEAR")
tr.font.size  = Pt(36)
tr.font.bold  = True
tr.font.color.rgb = DARK_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("Crisis-Led Evidence and Analysis for Response")
sr.font.size  = Pt(16)
sr.font.color.rgb = MID_BLUE
sr.font.bold  = False

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

org_p = doc.add_paragraph()
org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
org_p.add_run("Automated Situation Analysis — Project & Capabilities Report\n").font.size = Pt(12)
org_p.add_run("Media Monitoring & Analysis (MMP)\n").font.size = Pt(11)
date_run = org_p.add_run(f"April 2026")
date_run.font.size = Pt(11)
date_run.font.color.rgb = MID_BLUE

doc.add_page_break()


# ── 1. Executive Summary ──────────────────────────────────────────────────────
add_heading(doc, "1. Executive Summary")
add_divider(doc)
add_body(doc,
    "CLEAR (Crisis-Led Evidence and Analysis for Response) is an automated, end-to-end system "
    "that transforms raw humanitarian documents into a structured, interactive situation analysis "
    "dashboard aligned with the NRC Situation Analysis Framework (SAF). It addresses one of the "
    "most persistent bottlenecks in humanitarian response: the time and analytical capacity "
    "required to synthesise large volumes of information into actionable insights."
)
add_body(doc,
    "Where traditional situation analysis requires days to weeks of expert analyst time per "
    "country, CLEAR performs the same process in hours. It ingests reports, assessments, and "
    "situation reports from humanitarian data platforms, extracts and classifies analytical "
    "entries using purpose-trained NLP models, runs structured LLM-based analysis across the "
    "SAF pillar and sector framework, and produces an interactive, downloadable dashboard "
    "ready for operational use."
)
add_body(doc,
    "The system has been deployed for multi-country crisis analysis (Lebanon and Iran) and is "
    "designed to be extended to any country or crisis context. Alongside the core pipeline, the "
    "project has produced a suite of open-source NLP libraries covering document processing, "
    "entry extraction, humanitarian text classification, zero-shot classification, LLM inference "
    "orchestration, and geolocation extraction — all released under open licenses and available "
    "to the broader humanitarian data community."
)


# ── 2. Problem Statement ──────────────────────────────────────────────────────
add_heading(doc, "2. Problem Statement")
add_divider(doc)
add_body(doc,
    "Humanitarian situation analysis is a critical but resource-intensive process. Analysts must "
    "read and synthesise hundreds of documents, categorise information across structured "
    "analytical frameworks, extract key figures and risks, and produce coherent reports under "
    "significant time pressure. The current state-of-practice is characterised by several "
    "structural limitations:"
)
for bullet in [
    "Turnaround time: producing a situation analysis for a single country typically takes days to weeks of expert analyst effort.",
    "Scalability: the manual process cannot scale across multiple simultaneous crises without significant additional staffing.",
    "Currency: analyses quickly become outdated as new documents are published; updating them requires repeating most of the process.",
    "Consistency: different analysts apply frameworks differently, introducing variability in outputs across contexts.",
    "Accessibility: final outputs are typically locked in static PDF reports, limiting interactivity and reuse.",
]:
    add_bullet(doc, bullet)
add_body(doc,
    "CLEAR was designed to address all of these limitations through automation, without replacing "
    "the analytical judgment of humanitarian professionals.",
    space_after=12
)


# ── 3. System Overview ────────────────────────────────────────────────────────
add_heading(doc, "3. System Overview")
add_divider(doc)
add_body(doc,
    "The CLEAR system operates as a two-stage pipeline. The first stage ingests raw documents "
    "and produces structured analytical outputs. The second stage transforms those outputs into "
    "a fully interactive dashboard. An optional context generation module supplements the "
    "document-based analysis with live web-sourced background information."
)

add_heading(doc, "3.1  Stage 1 — From Documents to Structured Analysis", level=2)
add_body(doc,
    "The system begins by automatically collecting the most relevant humanitarian documents for "
    "the country in question — reports, situation reports, and assessments published on "
    "ReliefWeb. It then reads through all of them."
)
add_body(doc,
    "Rather than processing each document as a whole, the system breaks every document into "
    "short, focused excerpts — typically a few sentences each — where each excerpt deals with "
    "one distinct observation, figure, or finding. This is important because it allows the "
    "system to sort and analyse information at a much more precise level than the document as "
    "a whole would allow."
)
add_body(doc,
    "Each excerpt is then automatically labelled according to the SAF analytical framework. "
    "The system identifies what the excerpt is about — which type of crisis driver, which "
    "humanitarian condition, and which sector it relates to — across three dimensions "
    "simultaneously:"
)
styled_table(doc,
    headers=["Dimension", "What it captures"],
    rows=[
        ["Crisis drivers and context\n(Pillars 1D)", "What is causing displacement or crisis — including push factors, hazards, underlying vulnerabilities, and barriers to humanitarian access"],
        ["Humanitarian conditions\n(Pillars 2D)", "What impact the crisis is having on people — covering direct impact, humanitarian conditions on the ground, and populations at risk"],
        ["Sectors", "Which area of life is affected — Agriculture, Education, Food Security, Health, Livelihoods, Logistics, Nutrition, Protection, Shelter, or WASH"],
    ],
    col_widths=[1.8, 4.3]
)
add_body(doc,
    "Once excerpts have been labelled, the system conducts three parallel analyses for each "
    "country:"
)
add_body(doc, "Narrative analysis", bold=True, space_after=2)
add_body(doc,
    "For each analytical theme (a combination of crisis driver and sector), the system reads "
    "the most relevant excerpts and produces a structured summary. This includes: a written "
    "analytical answer, a list of identified risks with severity scores, key headline figures, "
    "the most pressing humanitarian needs, and recommended priority interventions — all with "
    "references back to the source documents.",
    space_after=8
)
add_body(doc, "Figures and statistics extraction", bold=True, space_after=2)
add_body(doc,
    "Separately, the system goes through every excerpt that mentions numbers — people displaced, "
    "killed, injured, or otherwise affected — and extracts each figure with its full context: "
    "what it refers to, when it was recorded, and where. This produces a clean, structured "
    "database of humanitarian statistics drawn directly from the source documents.",
    space_after=8
)
add_body(doc, "Country context", bold=True, space_after=2)
add_body(doc,
    "In parallel, the system uses live web search to compile a structured background brief on "
    "the country, covering nine areas: demographics, political situation, economy, "
    "socio-cultural context, security, legal and policy environment, infrastructure, "
    "environment, and humanitarian coordination. This context is displayed alongside the "
    "document-based analysis in the dashboard.",
    space_after=8
)

add_heading(doc, "3.2  Stage 2 — Dashboard Data and Filtering", level=2)
add_body(doc,
    "All analytical outputs are post-processed and filtered by severity before being loaded into "
    "the dashboard. Only the most critical information is surfaced, based on the following "
    "thresholds:"
)
styled_table(doc,
    headers=["Output", "Filter Criterion"],
    rows=[
        ["Sector severity matrix", "Maximum score per (Pillar 2D, Sector) pair — all scores included"],
        ["Humanitarian access barriers", "Risk score ≥ 8"],
        ["Critical sector numbers", "Risk score ≥ 8 and numeric value > 100"],
        ["Top sectoral needs", "Priority need score ≥ 9"],
        ["Top priority interventions", "Priority intervention score ≥ 9"],
        ["Displacement figures", "Risk score ≥ 9 and numeric value > 100"],
        ["Headline numbers (killed/injured)", "Risk score ≥ 8 and numeric value > 100"],
        ["Context risks", "Risk score ≥ 9"],
    ],
    col_widths=[2.8, 3.4]
)

add_heading(doc, "3.3  Interactive Dashboard", level=2)
add_body(doc,
    "The final output is a self-contained interactive HTML dashboard that requires no server or "
    "technical setup — it can be opened directly in any web browser or shared as a standalone "
    "file. The dashboard is organised into three tabs:"
)
add_body(doc, "Overview Tab", bold=True, space_after=2)
for item in [
    "Headline KPIs: killed, injured, and displaced figures",
    "Context risks by analytical pillar",
    "Hazards and threats vs. pre-crisis vulnerabilities",
    "Displacement drivers vs. return intentions",
]:
    add_bullet(doc, item)
add_body(doc, "Sectors Tab", bold=True, space_after=2)
for item in [
    "Critical numbers by sector",
    "Severity matrix across Impact, Humanitarian Conditions, and At Risk — per sector, with top-3 risks on hover",
    "Top needs by sector",
]:
    add_bullet(doc, item)
add_body(doc, "Response Planning Tab", bold=True, space_after=2)
for item in [
    "Humanitarian access barriers",
    "Priority interventions by sector",
]:
    add_bullet(doc, item)
add_body(doc,
    "The dashboard supports switching between countries (Lebanon and Iran in the current "
    "deployment) and includes a one-click PDF export that renders all three tabs into a "
    "print-ready document."
)


# ── 4. Analytical Framework ───────────────────────────────────────────────────
add_heading(doc, "4. Analytical Framework Alignment")
add_divider(doc)
add_body(doc,
    "CLEAR is structured around the NRC Situation Analysis Framework (SAF), which organises "
    "humanitarian analysis into a hierarchical set of pillars, subpillars, and sectors. This "
    "alignment ensures that all automated outputs map directly onto the analytical categories "
    "used by NRC analysts, making the dashboard immediately interpretable within existing "
    "workflows."
)
add_body(doc,
    "The classification taxonomy covers the three core SAF dimensions — Pillars 1D (causal and "
    "contextual factors), Pillars 2D (impact and conditions), and Sectors (thematic areas) — "
    "and enables cross-dimensional analysis, for example identifying which sectors are most "
    "affected by a specific type of shock, or which displacement drivers are linked to the "
    "most severe humanitarian conditions."
)
add_body(doc,
    "Every analytical output — risks, needs, interventions, numbers, and narrative answers — "
    "is anchored to a specific pillar/subpillar and sector combination, providing a structured, "
    "auditable chain from source document to analytical conclusion."
)


# ── 5. Open-Source Capabilities ───────────────────────────────────────────────
add_heading(doc, "5. Open-Source Capabilities Developed")
add_divider(doc)
add_body(doc,
    "Alongside the CLEAR pipeline, the project has produced six open-source Python libraries "
    "covering the full NLP stack required for humanitarian document analysis. All are released "
    "under the GNU Affero General Public License v3.0 (AGPL v3), publicly documented, and "
    "available as installable packages from GitHub. They are designed to be reusable "
    "independently of the CLEAR pipeline and to benefit the broader humanitarian data community."
)

# 5.1
add_heading(doc, "5.1  LLM Multiprocessing Inference", level=2)
add_body(doc,
    "A high-throughput batched inference library supporting OpenAI, Perplexity, and Ollama "
    "backends through a unified API. It handles asynchronous request batching, rate limiting, "
    "structured JSON parsing with fallback handling, and optional streaming output. In the "
    "CLEAR pipeline, it powers all parallel LLM calls in Stage 1 — the hundreds of concurrent "
    "analysis prompts that would otherwise need to be managed manually."
)
add_body(doc, "Key capabilities:", bold=True, space_after=2)
for item in [
    "Single function call handles batching, rate limiting, retries, and response parsing across all supported providers",
    "Structured output mode parses model JSON into Python objects with configurable fallback responses",
    "Streaming mode provides real-time output for local Ollama deployments",
    "Local Ollama models are pulled automatically if not present",
]:
    add_bullet(doc, item)
styled_table(doc,
    headers=["Provider", "Default Model", "Concurrency Mode", "API Key Required"],
    rows=[
        ["OpenAI", "gpt-4o-mini", "Async bulk, rate-limited", "Yes"],
        ["Perplexity", "llama-3.1-sonar-small-128k-chat", "Async bulk, rate-limited", "Yes"],
        ["Ollama (local)", "gemma3:4b-it-q4_K_M", "Sequential, memory-safe", "No"],
    ],
    col_widths=[1.2, 2.2, 2.0, 1.4]
)

# 5.2
add_heading(doc, "5.2  Documents Processing", level=2)
add_body(doc,
    "A Python library for extracting text, figures, tables, and metadata from PDF, DOCX, and "
    "PPTX documents using a combination of traditional extraction tools and vision-language "
    "models. It handles the full complexity of real-world humanitarian documents: scanned PDFs, "
    "mixed-content layouts, embedded charts and tables, and non-standard formatting."
)
add_body(doc, "Processing approach:", bold=True, space_after=2)
for item in [
    "Non-PDF formats are converted to PDF via LibreOffice",
    "Selectable text is extracted via PyMuPDF; scanned or image-based pages use Tesseract OCR",
    "Page images are analysed by YOLOv10 for layout detection, identifying figures and tables",
    "Detected figures and tables are captioned by a vision-language model (OpenAI or Ollama)",
    "Cover and last pages are passed to the VLM for structured metadata extraction (title, source organisation, date, document type, country)",
]:
    add_bullet(doc, item)
add_body(doc,
    "The library is available as a Python API, a command-line tool, a Flask REST API, and a "
    "Docker image with all dependencies pre-installed."
)

# 5.3
add_heading(doc, "5.3  Entry Extraction", level=2)
add_body(doc,
    "Extracts semantically coherent analytical entries from raw document text. Given the messy "
    "reality of humanitarian documents — OCR artifacts, missing punctuation, metadata headers, "
    "navigation text, mixed languages — this library performs a multi-step cleaning and "
    "segmentation pipeline before any classification or analysis is attempted."
)
add_body(doc, "Pipeline steps:", bold=True, space_after=2)
styled_table(doc,
    headers=["Step", "Method", "Purpose"],
    rows=[
        ["Punctuation restoration", "XLM-RoBERTa (ONNX)", "Fix sentence boundaries in OCR and PDF text"],
        ["Relevancy filtering", "SetFit model (sentence-level)", "Remove boilerplate, metadata, navigation, and non-substantive text"],
        ["Independence splitting", "SetFit model (topic-shift detection)", "Split groups of relevant sentences at topic boundaries"],
        ["Long entry chunking", "Sliding window", "Cap entries at a configurable maximum, with overlap for context continuity"],
        ["Page attribution", "Fuzzy sequence matching", "Map each entry back to its source page number"],
    ],
    col_widths=[2.0, 1.8, 2.8]
)
add_body(doc,
    "The result is a clean list of discrete, meaningful excerpts — each representing one "
    "analytical claim — with source page numbers attached. Boilerplate, URLs, headers, and "
    "all non-substantive content are silently discarded."
)

# 5.4
add_heading(doc, "5.4  Zero-Shot Classification", level=2)
add_body(doc,
    "A two-stage zero-shot classifier combining a multilingual NLI model for fast broad "
    "filtering with an optional LLM second pass for precise multi-label classification. "
    "The two-stage design significantly reduces LLM cost compared to pure LLM zero-shot "
    "approaches, while maintaining high accuracy."
)
add_body(doc, "Operational modes:", bold=True, space_after=2)
styled_table(doc,
    headers=["Mode", "Configuration", "Output"],
    rows=[
        ["Raw scoring", "First pass only, no threshold", "Confidence score per tag per text"],
        ["Threshold filtering", "First pass with score cutoff", "Tags above the threshold per text"],
        ["Full two-stage", "First pass + LLM refinement", "LLM-verified final tag list per text"],
    ],
    col_widths=[1.5, 2.2, 2.8]
)
add_body(doc,
    "A per-sentence mode is also available for long documents, where individual sentences "
    "carry more distinct signals than the full text — scores are aggregated by maximum across "
    "sentences. Supports CPU, CUDA, and Apple Silicon (MPS)."
)

# 5.5
add_heading(doc, "5.5  Humanitarian Extract Classificator (HumBERT)", level=2)
add_body(doc,
    "An inference library for classifying humanitarian text excerpts using HumBERT — a "
    "debiased transformer trained on HumSet, the largest publicly available humanitarian "
    "text classification dataset. HumBERT is directly used in the CLEAR pipeline to tag every "
    "extracted entry across the full SAF ontology."
)
add_body(doc,
    "The debiasing methodology addresses known biases in humanitarian text datasets (geographic, "
    "temporal, and source biases) and is documented in associated research. The trained model "
    "is downloaded automatically on first use."
)
add_body(doc, "Three-level classification hierarchy:", bold=True, space_after=2)
for item in [
    "Level 1 (HumBERT transformer): assigns SAF tags across Pillars 1D, Pillars 2D, and all ten sectors",
    "Level 2 (LLM): adds sub-category tags within each Level-1 label",
    "Problems (LLM): identifies specific humanitarian problems within each Level-2 category (e.g. Infrastructure damage, Food insecurity, Armed violence)",
]:
    add_bullet(doc, item)
add_body(doc,
    "A configurable precision/recall trade-off parameter allows operators to tune the "
    "classifier towards broader coverage (higher recall) or more confident tagging (higher "
    "precision) without retraining."
)

# 5.6
add_heading(doc, "5.6  Geolocations and Polygons Extraction", level=2)
add_body(doc,
    "Extracts geographical location mentions from multilingual text and resolves each mention "
    "to its administrative-level polygon data using the Fieldmaps dataset (ADM0–ADM4). "
    "The library handles multilingual input — location names in French, Arabic, or other "
    "languages are translated and matched against the reference database — and returns "
    "structured metadata including location IDs, names, and P-codes at each administrative "
    "level."
)
add_body(doc,
    "This capability enables geospatial analysis of humanitarian entries: mapping affected "
    "areas, aggregating data by administrative region, and linking textual reports to "
    "geographic boundaries for visualisation. It is available as a standalone library "
    "independently of the CLEAR pipeline."
)


# ── 6. Technology Overview ────────────────────────────────────────────────────
add_heading(doc, "6. Technology Overview")
add_divider(doc)
add_body(doc,
    "The following table summarises the key technologies employed across the CLEAR system "
    "and supporting libraries."
)
styled_table(doc,
    headers=["Function", "Technology"],
    rows=[
        ["Document scraping and ingestion", "ReliefWeb API"],
        ["PDF text extraction", "PyMuPDF, Tesseract OCR"],
        ["Figure and table extraction", "YOLOv10 layout detection + Vision-Language Model"],
        ["Office format conversion", "LibreOffice"],
        ["Punctuation restoration", "XLM-RoBERTa (ONNX runtime)"],
        ["Relevancy and independence models", "SetFit (sentence-transformers)"],
        ["Entry classification", "HumBERT (debiased transformer, trained on HumSet)"],
        ["Zero-shot classification", "bge-m3-zeroshot-v2.0 NLI + GPT-4o-mini"],
        ["Structured LLM analysis", "OpenAI GPT-4.1-mini"],
        ["Context and background generation", "OpenAI GPT-5-mini with live web search"],
        ["LLM inference orchestration", "LLM Multiprocessing Inference (async, rate-limited)"],
        ["Geolocation resolution", "Fieldmaps administrative polygon dataset (ADM0–ADM4)"],
        ["Interactive dashboard", "Self-contained HTML and JavaScript"],
    ],
    col_widths=[2.8, 3.5]
)


# ── 7. Deployment Status ──────────────────────────────────────────────────────
add_heading(doc, "7. Current Deployment Status")
add_divider(doc)
add_body(doc,
    "The CLEAR pipeline is currently stable and reproducible in a local environment. "
    "End-to-end outputs have been generated for Lebanon and Iran, with the interactive "
    "dashboard validated for operational use. The system is ready for the transition to "
    "a production deployment phase."
)
add_body(doc,
    "The six open-source libraries are publicly available, documented, and installable. "
    "They are actively used within the CLEAR pipeline and are available for independent "
    "adoption by other teams or organisations."
)
add_body(doc,
    "The primary gap between the current state and a fully operationalised service is "
    "infrastructure: the pipeline requires packaging as a set of orchestrated microservices "
    "with scheduling, persistent storage, monitoring, and access controls. This represents "
    "the critical path for the next phase."
)


# ── 8. Next Steps and Roadmap ─────────────────────────────────────────────────
add_heading(doc, "8. Next Steps and Delivery Roadmap")
add_divider(doc)
add_body(doc,
    "The following roadmap is structured around four delivery phases, ordered by dependency "
    "and strategic priority. Phase 1 establishes the production foundation required for all "
    "subsequent phases."
)

add_heading(doc, "8.1  Recommended Delivery Sequence", level=2)
for i, phase in enumerate([
    ("Phase 1 — Stabilise Production Foundation",
     "Implement orchestrated microservice architecture with scheduling, job monitoring, and a persistent classification database. Add authentication, role-based access controls, and secure secret management. Establish production-grade retry policies, run-state tracking, and observability. This phase is a prerequisite for all subsequent work."),
    ("Phase 2 — Scale Multilingual Quality",
     "Build AI-assisted gold datasets and run a multilingual benchmark campaign for entry extraction and classification. Retrain and upgrade models for non-Latin scripts and additional languages. Conduct analyst-in-the-loop validation sessions to tune risk cutoffs and scoring logic. Establish a versioned model registry and prompt audit trail."),
    ("Phase 3 — Optimise Performance and Coverage",
     "Refactor the document extraction flow to reduce unnecessary OCR and segmentation passes and optimise for volume. Integrate additional structured and unstructured data sources beyond ReliefWeb (including ACLED, OCHA, and social media where relevant). Define SLA targets and optimise parallelisation for high-volume country monitoring."),
    ("Phase 4 — Extend Advanced Outputs",
     "Design scenario templates and implement a forecasting workflow with evaluation criteria. Harden the processed-data retrieval chatbot for real analyst workflows, adding response validation and output guardrails. Improve the dashboard UX for non-technical users with clearer labels, guided explanations, and simplified views."),
], 1):
    add_body(doc, phase[0], bold=True, space_after=2)
    add_body(doc, phase[1], space_after=10)

add_heading(doc, "8.2  Priority Roadmap by Component", level=2)
add_body(doc,
    "The table below summarises the full set of identified improvements, their current "
    "status, and deployment readiness."
)
styled_table(doc,
    headers=["Component", "Current Status", "Recommended Actions", "Priority"],
    rows=[
        ["Pipeline Orchestration & Storage",
         "Runs locally",
         "Microservice architecture, scheduling, job monitoring, persistent classification DB",
         "Critical"],
        ["Entry Extraction",
         "Reliable for English; limited multilingual coverage",
         "Multilingual benchmark campaign; retrain for non-Latin scripts",
         "High"],
        ["Entry Classification (HumBERT)",
         "Functional on Latin languages (EN, FR, ES)",
         "Retrain on multilingual ground truth; extend language coverage",
         "High"],
        ["Core Analytical Outputs",
         "End-to-end outputs generated",
         "Analyst validation sessions; tune scoring logic; improve explainability",
         "High"],
        ["Numerical Extraction",
         "Partially implemented (displaced, killed, injured)",
         "Expand indicator catalog; improve temporal/location normalisation",
         "High"],
        ["Data Source Coverage",
         "Primarily ReliefWeb",
         "Integrate ACLED, OCHA, and additional connectors; add quality governance",
         "High"],
        ["Document Extraction",
         "Functional; runtime and cost high at volume",
         "Optimise extraction flow; expand native DOCX/PPTX support",
         "High"],
        ["Dashboard UX",
         "Suited to technical users",
         "Improve labels and navigation; add guided explanations for non-technical users",
         "High"],
        ["Idempotent Re-run & Recovery",
         "Supported locally",
         "Production-grade retry policies, run-state tracking, observability",
         "High"],
        ["Model & Prompt Versioning",
         "Scripts only",
         "Formal model registry, prompt versioning, audit trail",
         "High"],
        ["Security & Access Control",
         "No centralised controls",
         "Authentication, RBAC, secure secret management",
         "High"],
        ["Local LLM Processing",
         "Prototype exists (llama-cpp)",
         "Harden deployment profiles; simplify private on-prem setup",
         "Medium"],
        ["Geolocation Extraction",
         "Available with shapefile linkage",
         "Generalise adapters; improve confidence scoring",
         "Medium"],
        ["Forecasting & Scenarios",
         "Not implemented",
         "Design scenario templates; implement forecasting workflow",
         "Medium"],
        ["Processed-Data Chatbot",
         "Prototype exists",
         "Test on analyst workflows; add output schema guardrails",
         "Medium"],
        ["Cost & Latency Optimisation",
         "Acceptable at low volume",
         "Define SLA targets; optimise for high-volume country monitoring",
         "Medium"],
        ["Dashboard Packaging",
         "Manual static HTML/JS artifacts",
         "Standardised release process for client delivery lifecycle",
         "Medium"],
        ["Ground Truth & Benchmarking",
         "Ad hoc validation",
         "Versioned benchmark datasets; multilingual QA protocol",
         "High"],
    ],
    col_widths=[1.7, 1.5, 2.5, 0.8]
)


# ── 9. Licensing ──────────────────────────────────────────────────────────────
add_heading(doc, "9. Licensing and Open-Source Commitment")
add_divider(doc)
add_body(doc,
    "All six libraries developed as part of this project are released under the GNU Affero "
    "General Public License v3.0 (AGPL v3). This license:"
)
for item in [
    "Permits free use, modification, and distribution of the software",
    "Requires that any derivative works be released under the same license",
    "Extends copyleft obligations to network use: organisations running the software as a service must also open-source their modifications",
]:
    add_bullet(doc, item)
add_body(doc,
    "This commitment reflects the project's positioning within the humanitarian data ecosystem: "
    "the tools are designed to be shared, built upon, and improved by the community they serve. "
    "For questions or collaboration enquiries, contact: selimfek@gmail.com"
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/Users/sfekih/Documents/MMP/CLEAR-AutomatedAnalysis/CLEAR_Project_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
